"""
make_speech.py — 生成组会汇报演讲稿 Word,与 PPT 6 页一一对应
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/mnt/d/pan-caner/group_meeting_report/speech.docx"

doc = Document()

# Default style: Calibri 11
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Set East Asian font fallback for Chinese
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rPr.append(rFonts)
rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

# Set page margins
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ===== Helpers =====
def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    rPr2 = r._element.get_or_add_rPr()
    rFonts2 = OxmlElement("w:rFonts")
    rFonts2.set(qn("w:eastAsia"), "Microsoft YaHei")
    rPr2.append(rFonts2)

def add_page_header(page_num, title):
    """Add a colored page header like the PPT."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    # Page number badge as small text
    r1 = p.add_run(f"[P{page_num}] ")
    r1.font.size = Pt(14)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    # Title
    r2 = p.add_run(title)
    r2.font.size = Pt(16)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    rPr2 = r2._element.get_or_add_rPr()
    rFonts2 = OxmlElement("w:rFonts")
    rFonts2.set(qn("w:eastAsia"), "Microsoft YaHei")
    rPr2.append(rFonts2)

def add_meta(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    rPr2 = r._element.get_or_add_rPr()
    rFonts2 = OxmlElement("w:rFonts")
    rFonts2.set(qn("w:eastAsia"), "Microsoft YaHei")
    rPr2.append(rFonts2)

def add_speech(text):
    """Body paragraph — the speech script."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.line_spacing = 1.4
    r = p.add_run(text)
    r.font.size = Pt(11.5)
    r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    rPr2 = r._element.get_or_add_rPr()
    rFonts2 = OxmlElement("w:rFonts")
    rFonts2.set(qn("w:eastAsia"), "Microsoft YaHei")
    rPr2.append(rFonts2)

def add_tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run("🎤 演讲提示: " + text)
    r.font.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
    rPr2 = r._element.get_or_add_rPr()
    rFonts2 = OxmlElement("w:rFonts")
    rFonts2.set(qn("w:eastAsia"), "Microsoft YaHei")
    rPr2.append(rFonts2)

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("─" * 60)
    r.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)

# ===== Title page =====
add_title("肺癌 9 基因突变预测项目 — 组会演讲稿")

# Subtitle
p = doc.add_paragraph()
r = p.add_run("组会汇报 · 公开数据集调研与处理流程 · 6 页 PPT 对应讲解")
r.font.size = Pt(11)
r.font.italic = True
r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

# Meta info
add_meta("Reporter: zzy-ustb  |  Date: 2026-07-12  |  Duration: ~12 min")

add_divider()

# Intro
add_h2("开场 (1 min)")
add_speech(
    "各位老师同学好,今天汇报我近期在 9 基因肺癌预测项目上的数据调研与处理工作。"
    "这一阶段的核心目标是构建 patient-level 的 mutation 标签数据集,以及相应的 WSI 图像数据,"
    "为后续训练 pipeline 打下基础。今天主要讲四件事:9 基因 panel 的确定、MC3 mutation 数据的调研与提取、"
    "TCGA-LUAD WSI DICOM 数据的下载处理、以及 DeepGEM pipeline 的借鉴分析。"
)

# ===== Page 1 =====
add_page_header(1, "封面 — 项目总览")
add_meta("PPT 第 1 页 (Title Slide)")
add_h2("讲解要点")
add_speech(
    "这一页是封面,核心信息是我们要预测肺癌 9 个驱动基因的突变状态。"
    "三个关键数字先点一下:421 个 TCGA-LUAD 病例带 mutation 标签,"
    "其中 50 个已经下载了 WSI 数据,模型最终的目标是给定一张 H&E 病理切片,"
    "预测 11 个基因的突变概率。"
)
add_speech(
    "项目背景:这是与聚时医疗的横向课题,我们的硬指标是准确率 95% 以上、灵敏度 90% 以上,"
    "技术栈涉及 DINOv3 自监督预训练、三维病理,以及和 LIS-PACS 接口的对接。"
    "今天汇报的是 9 基因 panel 的数据基础部分。"
)
add_tip("开场用 30 秒,重点把 9 基因项目和今天要讲的 4 件事交代清楚即可。")

# ===== Page 2 =====
add_page_header(2, "项目背景与 9 基因 Panel")
add_meta("PPT 第 2 页 (Project & Gene Panel)")
add_h2("讲解要点")
add_speech(
    "这一页讲两件事:项目背景,以及 9 基因是怎么来的。"
)
add_speech(
    "项目背景简单提一下:合作方是北科大和聚时医疗,合同有三个交付节点,准确率 95% 是硬指标。"
    "数据集策略上,内部数据来自聚时,外部公开数据我们用 TCGA-LUAD 作为训练基底,"
    "同时为了审稿人认账,异源数据是必须的(钟团队的 CJFH 数据后续会接)。"
)
add_speech(
    "9 基因 panel 是这样定的:DeepGEM 已经验证过的 6 个基因(EGFR、KRAS、ALK、ROS1、TP53、BRAF)是基础,"
    "加上 NCCN 指南和合同附件里强调的商业靶向药相关基因(PIK3CA、ERBB2/HER2、NRAS、RET),"
    "实际训练时我们冗余到 11 个基因,留一些灵活度。"
)
add_speech(
    "阳性定义比较直接:只要 MC3 v0.2.8 PUBLIC MAF 里这个 case 有至少一个 PASS 的功能改变性突变,"
    "我们就把这个基因标记为 1。所谓功能改变性,指的是 Missense、Nonsense、Frameshift、Splice 这 9 类,"
    "同义突变不算。"
)
add_tip("重点是让听众理解 panel 来源的合理性,以及为什么有冗余。可以略过 11/9 的细节争论。")

# ===== Page 3 =====
add_page_header(3, "MC3 MAF 数据调研与提取")
add_meta("PPT 第 3 页 (Mutation Labels)")
add_h2("讲解要点")
add_speech(
    "这一页是核心数据工作之一:TCGA-LUAD mutation 标签的提取。数据源是 MC3 —— "
    "TCGA 联合 7 个变异 caller 做的 pan-cancer consensus 突变集合。"
)
add_speech(
    "调研过程有几点值得讲。第一,MC3 没有 HTTP 直链,走的是 Synapse,需要注册账号 + PAT 鉴权。"
    "第二,GDC Legacy Archive 之前能下 MAF,但已经关闭了,新 GDC API 暂时也不返回完整 slide image。"
    "我们最终的方案是 Synapse syn7824274,这个数据集是 PUBLIC 的,不需要 dbGaP 授权,只需要 Synapse 账号。"
)
add_speech(
    "提取流程的关键技术点有三个。第一,我们用 streaming gzip 读取 753 MB 的 MAF 文件,内存占用只是单行,"
    "不会爆。第二,cBioPortal 提供了 584 个 TCGA-LUAD 候选 case_id 的白名单 —— 这步很关键,"
    "因为我们最初的 is_luad() 函数只认 barcode 第 2 段等于 '05',只覆盖了 30 个 cases,严重不全。"
    "TCGA 的 barcode 第 2 段其实是 site code,不是 disease code,LUAD 散落在 22 个 site 上。"
    "改用 cBioPortal 白名单后,最终匹配到 421 个 case,完整覆盖。"
    "第三,我们做了 4 层过滤:case_id 白名单、PASS 过滤、11 基因白名单、9 类功能改变变异类型,"
    "最后输出的是 patient-level 的 0/1 矩阵,加上一列 any_9gene_positive 和 n_genes_mutated 做汇总。"
)
add_speech(
    "右边的图是 11 基因的突变 prevalence,排序后可以直观看到 TP53 最高(61.5%)、KRAS(35.9%)、EGFR(15.9%) 这些,"
    "和文献报告的范围基本一致,MC3 因为用了多 caller 并集,数值略偏高是可预期的。"
)
add_tip("重点让听众明白:我们不是简单的'下一个 CSV',而是碰到了 GDC 迁移 + barcode 误读两个工程坑并解决了。")

# ===== Page 4 =====
add_page_header(4, "TCGA-LUAD WSI DICOM 数据下载与处理")
add_meta("PPT 第 4 页 (DICOM WSI)")
add_h2("讲解要点")
add_speech(
    "这一页讲 WSI 图像数据。WSI 是 Whole Slide Image,一张病理切片的全分辨率扫描,通常是几千乘几千像素。"
)
add_speech(
    "格式上,业界有三家主流:.svs 是 Leica/Aperio 的,.sdpc 是 Philips 的,两者都能被 openslide 直接读取。"
    "但 TCGA 数据现在通过 IDC 分发的都是 .dcm 格式 —— DICOM 标准,跨厂商。"
    "DICOM 的好处是元数据标准化,坏处是 openslide 不支持,需要用 wsidicom 这个专门的库。"
)
add_speech(
    "DICOM 内部结构我做了一个 demo 给大家看,关键 tag 有 PatientID —— 这是和我们 mutation CSV 的 join key;"
    "Modality 区分 SM、SEG、ANN、CT 几种,SM 是 Slide Microscopy 也就是真切片,是我们要用的;"
    "还有 TotalPixelMatrix、RowPosition 这些是拼接用的坐标信息。"
)
add_speech(
    "关于拼接:DICOM WSI 把一张大切片切成几十到几百个 tile,每个 tile 是独立 .dcm 文件。"
    "要还原完整切片,需要按 RowPositionInTotalImagePixelMatrix 和 ColumnPositionInTotalImagePixelMatrix 这两个 tag 拼,"
    "这是 DICOM Part 3 强制要求的 tag,所有厂商(Philips、Leica、Hamamatsu)都填,所以不会拼错。"
)
add_speech(
    "下方是数据集 stat tile:50 个 case 已经下载到本地,平均每张切片 9.6×7.5 mm,物镜 20 倍。"
    "我们这个阶段够做 prototype 了,扩到 200-300 cases 是后续工作。"
)
add_tip("这一页技术细节多,可以快讲。把 DICOM 是真切片、20× 物镜、4 类 modality 讲清楚就行。")

# ===== Page 5 =====
add_page_header(5, "真实切片可视化")
add_meta("PPT 第 5 页 (Real WSI Visualization)")
add_h2("讲解要点")
add_speech(
    "这一页是科研 figure,给大家看一张真实的肺癌切片。"
    "左边是从 TCGA-05-4245 这个 case 中央 1024×1020 像素的 patch,右边是整张切片的缩略图。"
)
add_speech(
    "我们看左边的 patch:H&E 染色 —— 紫色是苏木精染的细胞核,粉红色是伊红染的胞质和间质,"
    "中间那些白色空洞是肺泡结构。这张切片能看到肿瘤细胞密度高、核异型明显,"
    "是典型的肺腺癌 LUAD 表型。"
)
add_speech(
    "右边的关键事实是从 DICOM header 直接读出来的:PatientID 是 TCGA-05-4245,"
    "Protocol ID 是 TCGA-LUAD,制造商是 Leica Biosystems,原始 scanner 是 Aperio(后用 PixelMed 工具转成了 DICOM)。"
    "Series Description 是 'Frozen HE TP BS1' —— Frozen 表示冷冻切片,HE 就是 H&E 染色,TP 是 Topo pathologist,BS1 是 block section 1。"
    "Container ID 是 TCGA-05-4245-01A-01-BS1,这个 ID 直接对应到 LIS 系统里的玻片实物。"
)
add_speech(
    "这张图验证了 DICOM 到像素的通路是通的,后续可以直接进 DINOv3 backbone 做训练。"
)
add_tip("这是这张片子最'感性'的部分,可以多停留 30 秒让大家看清组织结构。")

# ===== Page 6 =====
add_page_header(6, "DeepGEM Pipeline 借鉴与下一步计划")
add_meta("PPT 第 6 页 (DeepGEM & Next Steps)")
add_h2("讲解要点")
add_speech(
    "最后一页讲两件事:DeepGEM 的 pipeline 设计,以及我们的下一步计划。"
)
add_speech(
    "DeepGEM 是腾讯 AI Lab 发表在 Lancet Oncology 的工作,核心思路是 self-supervised learning + 多实例学习 + label disambiguation,"
    "在 TCGA 和 7 个外部中心验证了 EGFR 等基因的预测。"
)
add_speech(
    "它的 pipeline 设计是 4 步:第一步 WSI cropping,把整张切片切成 1120×1120 像素的 patch,在 20 倍物镜下;"
    "第二步去掉背景 patch(没组织的);第三步用 CTransPath 或 EfficientNet-B0 提取 patch 特征,存成 .pkl;"
    "第四步合并每个 WSI 的所有 patch 特征,作为一个 bag,丢进 MIL 模型做训练。"
)
add_speech(
    "关键的洞察是:DeepGEM 真正训练的输入不是 WSI 本身,而是 patch features 的 .pkl。"
    "这对我们很有利 —— 我们目前 DICOM 的物镜倍数是 20 倍,跟 DeepGEM 完全一致,"
    "意味着 DICOM 不需要任何下采样就能直接走它的 pipeline。"
    "我们只需要写一个薄薄的 adapter,让 wsidicom 模拟 openslide 的 read_region 接口,"
    "DeepGEM 的 step1 到 step4 一行都不用改。"
)
add_speech(
    "下一步计划是这样的:第一步写 DICOM adapter,预计 30 分钟搞定;"
    "第二步对 50 个 case 跑 patch 提取,大约 5000 个 patch,落盘 PNG;"
    "第三步用 CTransPath 提特征,存 combined_feat.pkl;"
    "第四步用 DeepGEM 的架构微调,以 MC3 9-gene label 为训练目标,冲合同 95% 的准确率指标;"
    "第五步异源验证,接钟团队的 CJFH 数据,这块需要 dbGaP 授权,流程会比较长。"
)
add_tip("结尾给听众一个清晰的下一步路线图,以及时间预期。可以提一句:'如果一切顺利,本周内能完成 patch 提取的 POC。'")

add_divider()

# ===== 结尾 =====
add_h2("结束语 (30 sec)")
add_speech(
    "以上就是近期工作的汇报。总结一下:mutation 标签这边,421 个 TCGA-LUAD cases × 11 基因的 patient-level 数据已经完整落地;"
    "WSI 数据这边,50 个 case 已经下载并验证可以读取和拼接;pipeline 这边,DeepGEM 的方案和我们数据兼容性很高,"
    "下一步可以直接对接。"
)
add_speech(
    "请各位老师同学批评指正。"
)

doc.save(OUT)
print(f"演讲稿 saved: {OUT}")
import os
print(f"Size: {os.path.getsize(OUT)/1024:.1f} KB")